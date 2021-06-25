from docx import Document
from docx.shared import Inches

# GLOBAL VARIABLES
document = Document()

# FUNCTIONS
def user_image():
    print('Choose your photo:')
    
    path = input('Enter the image path > ')
    
    document.add_picture(path, width=Inches(1.25))
    
    print('\n')
    
    personal_data()


def personal_data():
    name = input('What is your name? > ')
    phone_number = input('What is your phone number? > ')
    email = input('What is your email? > ')
    place = input('Where do you live? > ')
    about_me = input('Tell me about you > ')
    
    document.add_paragraph(f'{name}  |  {phone_number}  |  {email}  | {place}')
    
    document.add_heading('About me')
    document.add_paragraph(about_me)
    
    print('\n')
    
    education()
    
def education():
    i = 1
    
    document.add_heading('Education')
    paragraph = document.add_paragraph()
    
    school_number = int(input('How many schools would you like to add? > '))
    
    while i <= school_number:
        school = input('School name > ')
        form_date = input('Fom date > ')
        to_date = input('To date > ')
        education_lvl = input('Elementary School, Middle School, High School, University? > ')
        degree = input('Degree or any award? (Leave blank if you dont want to add something here) > ')
        
        if degree:
            paragraph.add_run(f'{school} ').bold = True
            paragraph.add_run(f'{form_date} - {to_date} ').italic = True
            paragraph.add_run(f'({education_lvl})\n')
            paragraph.add_run(f'{degree}\n')
        else:
            paragraph.add_run(f'{school} ').bold = True
            paragraph.add_run(f'{form_date} - {to_date} ').italic = True
            paragraph.add_run(f'({education_lvl})\n')
        
        i += 1
    
    print('\n')
    
    experience()
    
    
def experience():
    i = 1
    
    document.add_heading('Work Experience')
    paragraph = document.add_paragraph()
    
    company_number = int(input('How many companies would you like to add? > '))
    
    while i <= company_number:
        company = input('Company name > ')
        from_date = input('From date > ')
        to_date = input('ot date > ')
        company_exp = input('Tell me about your experience at the company > ')
        
        paragraph.add_run(f'{company} ').bold = True
        paragraph.add_run(f'{from_date} - {to_date}\n').italic = True
        paragraph.add_run(company_exp + '\n')
        
        i += 1
    
    print('\n')
    
    skills()

def skills():
    document.add_heading('Skills')
    
    skill = input('Skill name > ')
    p = document.add_paragraph(skill)
    p.style = 'List Bullet'
    
    while True:
        more_skills = input('Do you have more skills? Yes or No > ')
        if more_skills.casefold() == 'yes':
            skill = input('Skill name > ')
            p = document.add_paragraph(skill)
            p.style = 'List Bullet'
        else:
            break
          

def save():
    document.save('cv.docx')
    
# OUTPUT
print('Welcome to CV Maker!')

user_image()
save()